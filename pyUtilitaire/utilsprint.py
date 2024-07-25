from docx import Document




###############################################################################
# Sortie console  
###############################################################################

def printentete():
    print("="*80)
    print(f'{"DESIGNATION":45} {"SYMBOLE":<10} {"VALEUR":>10} {"UNITE":>11}')
    print("="*80)

def printligne(designation, symbole, unite, valeur):
    print(f'    {designation:41} {symbole:<10} {valeur:>10} {unite:>11}')
        
def printsep():
    print("-"*80)

def printfintab():
    print("="*80)

def printverification():
    print("Vérifications :")
    print("---------------")

def printVERIFIE(message):
    print("    VERIFIE : ", message)

def printNONVERIFIE(message):
    print("NON VERIFIE : ", message)
    
def printMESSAGE(message):
    print("              ", message)  
    
###############################################################################
# Sortie fichier txt
###############################################################################
 
def printfile(texte, f):
    print(texte, file=f)

def printentetefile(f):
    print("="*77, file=f)
    print(f'{"DESIGNATION":45} {"SYMBOLE":<10} {"VALEUR":^10} {" "} {"UNITE":^10}', file=f)
    printfile("="*77, f)

def printlignefile(designation, symbole, unite, valeur, file):
    print(f'{designation:45} {symbole:<15} {valeur:>10} {"   "} {unite:<10}', file=file)

def printsepfile(f):
    print("-"*77, file=f)


###############################################################################
# Sortie fichier docx
###############################################################################

def print_entete_docx(doc):
    doc.add_paragraph(f'{"DESIGNATION":45} {"SYMBOLE":<10} {"VALEUR":<10} {"UNITE":<10}')
    doc.add_paragraph("-"*75)
    
def print_ligne_docx(doc, designation, symbole, unite, valeur):
    doc.add_paragraph(f'    {designation:41} {symbole:<10} {valeur:<10} {unite:<10}')


###############################################################################
# Test
###############################################################################
    
if __name__ == "__main__":
    largeur = 16 / 100
    hauteur = 13 / 100
    longueur = 2100 / 100
    
    printentete()   
    print("Dimensions de l'élement")
    printligne("Largeur", "la", "cm", f"{largeur*100:.1f}")
    printligne("Hauteur", "h", "cm", f"{hauteur*100:.1f}")
    printligne("Longueur", "L", "cm", f"{longueur*100:.1f}")

